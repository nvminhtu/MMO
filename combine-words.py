from docxcompose.composer import Composer
from docx import Document
import glob

doc1 = Document("doc1.docx")
composer = Composer(doc1)

for f in glob.glob("*.docx"):
    composer.append(Document(f))
# doc2 = Document("doc2.docx")

composer.save("abc-done.docx")